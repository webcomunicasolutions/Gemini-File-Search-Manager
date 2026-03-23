from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from google import genai
from google.genai import types
import os
import time
import json
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import logging
import mimetypes
import requests as http_requests
from docx import Document as DocxDocument
from openpyxl import load_workbook

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app, origins=['http://localhost:5001', 'http://127.0.0.1:5001'])

# Configure logging - Configuración de logs
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration constants - Constantes de configuración
UPLOAD_FOLDER = 'uploads'
# Expanded allowed extensions to match all File Search supported formats
ALLOWED_EXTENSIONS = {
    # Documents
    'txt', 'pdf', 'doc', 'docx', 'odt', 'rtf', 'md', 'markdown',
    # Spreadsheets
    'csv', 'tsv', 'xlsx', 'xls', 'xlsm', 'xlsb', 'ods',
    # Presentations
    'pptx', 'ppt', 'odp',
    # Data formats
    'json', 'xml', 'yaml', 'yml', 'sql',
    # Code files
    'py', 'js', 'jsx', 'ts', 'tsx', 'java', 'c', 'cpp', 'h', 'hpp',
    'cs', 'go', 'rs', 'php', 'rb', 'swift', 'kt', 'scala', 'pl',
    'r', 'hs', 'erl', 'lisp', 'lua', 'sh', 'bash', 'zsh', 'dart',
    # Web
    'html', 'htm', 'css', 'scss', 'sass',
    # Scientific
    'ipynb', 'bib', 'bibtex', 'tex',
    # Archives
    'zip'
}
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
MAX_HISTORY = 7  # Conversation history limit - Límite de historial de conversación

# Complete MIME type mapping for Gemini File Search API
# Based on official documentation: https://ai.google.dev/gemini-api/docs/file-search
MIME_TYPE_MAPPING = {
    # Documents
    'pdf': 'application/pdf',
    'doc': 'application/msword',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'odt': 'application/vnd.oasis.opendocument.text',
    'rtf': 'text/rtf',
    'txt': 'text/plain',
    'md': 'text/markdown',
    'markdown': 'text/markdown',

    # Spreadsheets (CRITICAL - These were failing without explicit MIME types)
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'xls': 'application/vnd.ms-excel',
    'xlsm': 'application/vnd.ms-excel.sheet.macroEnabled.12',
    'xlsb': 'application/vnd.ms-excel.sheet.binary.macroEnabled.12',
    'csv': 'text/csv',
    'tsv': 'text/tab-separated-values',
    'ods': 'application/vnd.oasis.opendocument.spreadsheet',

    # Presentations
    'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
    'ppt': 'application/vnd.ms-powerpoint',
    'odp': 'application/vnd.oasis.opendocument.presentation',

    # Data formats
    'json': 'application/json',
    'xml': 'application/xml',
    'yaml': 'text/yaml',
    'yml': 'text/yaml',
    'sql': 'application/sql',

    # Programming Languages
    'py': 'text/x-python',
    'js': 'text/javascript',
    'jsx': 'text/jsx',
    'ts': 'application/typescript',
    'tsx': 'text/tsx',
    'java': 'text/x-java',
    'c': 'text/x-c',
    'cpp': 'text/x-c++src',
    'cc': 'text/x-c++src',
    'cxx': 'text/x-c++src',
    'h': 'text/x-chdr',
    'hpp': 'text/x-c++hdr',
    'cs': 'text/x-csharp',
    'go': 'text/x-go',
    'rs': 'text/x-rust',
    'php': 'application/x-php',
    'rb': 'text/x-ruby-script',
    'swift': 'text/x-swift',
    'kt': 'text/x-kotlin',
    'scala': 'text/x-scala',
    'pl': 'text/x-perl',
    'r': 'text/x-rsrc',
    'hs': 'text/x-haskell',
    'erl': 'text/x-erlang',
    'lisp': 'text/x-lisp',
    'lua': 'text/x-lua',
    'sh': 'application/x-sh',
    'bash': 'application/x-sh',
    'zsh': 'application/x-zsh',
    'dart': 'application/vnd.dart',

    # Web Technologies
    'html': 'text/html',
    'htm': 'text/html',
    'css': 'text/css',
    'scss': 'text/x-scss',
    'sass': 'text/x-sass',

    # Scientific & Academic
    'ipynb': 'application/vnd.jupyter',
    'bib': 'text/x-bibtex',
    'bibtex': 'text/x-bibtex',
    'tex': 'application/x-tex',

    # Archives
    'zip': 'application/zip',

    # Other
    'ics': 'text/calendar',
    'vcard': 'text/vcard',
    'vcf': 'text/vcard'
}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize Gemini client - API key del usuario
api_key = os.getenv('GEMINI_API_KEY')
if not api_key:
    raise ValueError("GEMINI_API_KEY not found in environment variables. Please set GEMINI_API_KEY in your .env file")

client = genai.Client(api_key=api_key)

# Global state management - Gestión de estado global
conversation_history = []
file_search_store = None
uploaded_files = []  # Track uploaded files with metadata
PERSISTENCE_FILE = 'store_state.json'

# ============================================
# STATE PERSISTENCE FUNCTIONS - Funciones de persistencia de estado
# ============================================

def load_state():
    """Load persisted state from JSON file on startup - Cargar estado persistido desde archivo JSON"""
    global file_search_store, uploaded_files
    try:
        if os.path.exists(PERSISTENCE_FILE):
            with open(PERSISTENCE_FILE, 'r') as f:
                state = json.load(f)
                store_name = state.get('store_name')
                uploaded_files = state.get('uploaded_files', [])

                if store_name:
                    # Verify the store still exists
                    try:
                        file_search_store = client.file_search_stores.get(name=store_name)
                        logger.info(f"Restored file search store: {store_name} with {len(uploaded_files)} files")
                    except Exception as e:
                        logger.warning(f"Stored file search store not found: {e}")
                        file_search_store = None
                        uploaded_files = []
    except Exception as e:
        logger.error(f"Error loading state: {e}")

def save_state():
    """Save current state to JSON file - Guardar estado actual a archivo JSON"""
    try:
        state = {
            'store_name': file_search_store.name if file_search_store else None,
            'uploaded_files': uploaded_files
        }
        with open(PERSISTENCE_FILE, 'w') as f:
            json.dump(state, f, indent=2)
        logger.info("State saved successfully")
    except Exception as e:
        logger.error(f"Error saving state: {e}")

def load_state_value(key, default=None):
    """Load a specific value from state file"""
    try:
        if os.path.exists(PERSISTENCE_FILE):
            with open(PERSISTENCE_FILE, 'r') as f:
                state = json.load(f)
                return state.get(key, default)
    except Exception:
        pass
    return default

def save_state_value(key, value):
    """Save a specific key-value to state file"""
    try:
        state = {}
        if os.path.exists(PERSISTENCE_FILE):
            with open(PERSISTENCE_FILE, 'r') as f:
                state = json.load(f)
        state[key] = value
        with open(PERSISTENCE_FILE, 'w') as f:
            json.dump(state, f, indent=2)
    except Exception as e:
        logger.error(f"Error saving state value {key}: {e}")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ============================================
# METADATA HELPER - Extracción y fusión de metadatos
# ============================================

def extract_document_metadata(doc, uploaded_files):
    """Extract and merge Gemini + local metadata for a document.

    Gemini metadata is fetched from the API response. Local metadata
    from uploaded_files takes precedence (allows editing without re-upload).

    Args:
        doc: Document resource object from Gemini API
        uploaded_files: List of locally tracked file dicts

    Returns:
        dict: Merged metadata with local values overriding Gemini values
    """
    custom_metadata = {}
    if hasattr(doc, 'custom_metadata') and doc.custom_metadata:
        for metadata in doc.custom_metadata:
            key = getattr(metadata, 'key', '')
            if hasattr(metadata, 'string_value'):
                custom_metadata[key] = getattr(metadata, 'string_value', '')
            elif hasattr(metadata, 'numeric_value'):
                custom_metadata[key] = getattr(metadata, 'numeric_value', 0)

    # Merge with local metadata - local overrides Gemini
    for file_info in uploaded_files:
        if file_info.get('document_id') == doc.name:
            local_metadata = file_info.get('custom_metadata', {})
            custom_metadata.update(local_metadata)
            break

    return custom_metadata

def get_mime_type(filename):
    """
    Detect MIME type from file extension with fallback.

    Args:
        filename (str): Name of the file

    Returns:
        str: MIME type string
    """
    # Get file extension
    file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    # Priority 1: Check manual mapping (most reliable for File Search API)
    if file_ext in MIME_TYPE_MAPPING:
        return MIME_TYPE_MAPPING[file_ext]

    # Priority 2: Fallback to mimetypes module
    mime_type, _ = mimetypes.guess_type(filename)
    if mime_type:
        return mime_type

    # Priority 3: Safe default
    logger.warning(f"Could not determine MIME type for {filename}, using default")
    return 'application/octet-stream'

def extract_text_from_docx(file_path):
    """
    Extract text content from a DOCX file.

    Args:
        file_path (str): Path to the DOCX file

    Returns:
        str: Extracted text content
    """
    try:
        doc = DocxDocument(file_path)
        paragraphs = []

        # Extract all paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        text = '\n'.join(paragraphs)
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_xlsx(file_path):
    """
    Extract text content from an XLSX file.

    Args:
        file_path (str): Path to the XLSX file

    Returns:
        str: Extracted text content
    """
    wb = None
    try:
        wb = load_workbook(filename=file_path, read_only=True, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")

            # Read up to 100 rows to avoid overwhelming the model
            max_rows = min(sheet.max_row, 100) if sheet.max_row else 100

            for row in sheet.iter_rows(min_row=1, max_row=max_rows, values_only=True):
                # Convert row values to strings, skip empty rows
                row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)

        text = '\n'.join(text_parts)
        logger.info(f"Extracted {len(text)} characters from XLSX ({len(wb.sheetnames)} sheets)")
        return text

    except Exception as e:
        logger.error(f"Error extracting text from XLSX: {str(e)}")
        return ""
    finally:
        # IMPORTANT: Always close the workbook to release file handle
        if wb is not None:
            try:
                wb.close()
                logger.info("XLSX workbook closed successfully")
            except Exception as close_error:
                logger.warning(f"Error closing workbook: {close_error}")

# ============================================
# MAIN ROUTES - Rutas principales
# ============================================

@app.route('/')
def index():
    return render_template('index.html')

# ============================================
# FILE UPLOAD WITH PERSISTENCE - Carga de archivos con persistencia
# Timeout de 120 segundos configurado
# ============================================

@app.route('/upload', methods=['POST'])
def upload_file():
    global file_search_store, uploaded_files

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not supported'}), 400

    filepath = None

    try:
        # Save file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Get file size
        file_size = os.path.getsize(filepath)
        logger.info(f"File saved: {filename}, size: {file_size} bytes")

        # Get custom metadata and chunking config from request
        metadata_json = request.form.get('metadata', '{}')
        chunking_json = request.form.get('chunking_config', '{}')
        store_name_param = request.form.get('store_name', '')
        custom_metadata = json.loads(metadata_json)
        chunking_config = json.loads(chunking_json)

        # Determine which store to use
        target_store = None
        if store_name_param:
            # Use the specified store
            logger.info(f"Using specified store: {store_name_param}")
            target_store_name = store_name_param
        elif file_search_store is None:
            # Create new default store if none exists
            logger.info("Creating new file search store")
            file_search_store = client.file_search_stores.create(
                config={'display_name': 'RAG-App-Store'}
            )
            target_store_name = file_search_store.name
        else:
            # Use existing default store
            target_store_name = file_search_store.name

        # Detect MIME type explicitly (CRITICAL for CSV, XLSX, etc.)
        mime_type = get_mime_type(filename)
        logger.info(f"Detected MIME type: {mime_type} for {filename}")

        # Build upload config WITHOUT MIME type (will be auto-detected)
        upload_config = {
            'display_name': filename
        }

        # Add custom metadata if provided
        if custom_metadata:
            metadata_list = []
            for key, value in custom_metadata.items():
                if isinstance(value, (int, float)):
                    metadata_list.append({"key": key, "numeric_value": value})
                else:
                    metadata_list.append({"key": key, "string_value": str(value)})
            upload_config['custom_metadata'] = metadata_list

        # Add chunking config if provided
        if chunking_config and chunking_config.get('enabled'):
            upload_config['chunking_config'] = {
                'white_space_config': {
                    'max_tokens_per_chunk': chunking_config.get('max_tokens_per_chunk', 200),
                    'max_overlap_tokens': chunking_config.get('max_overlap_tokens', 20)
                }
            }

        # TRY DIRECT UPLOAD FIRST (upload_to_file_search_store)
        uploaded_api_file = None
        try:
            logger.info(f"Attempting direct upload for {filename} (MIME type will be auto-detected)")
            operation = client.file_search_stores.upload_to_file_search_store(
                file=filepath,
                file_search_store_name=target_store_name,
                config=upload_config
            )
        except Exception as upload_error:
            # FALLBACK: Use Files API + importFile for problematic files (CSV, large files, etc.)
            logger.warning(f"Direct upload failed: {upload_error}")
            logger.info(f"Falling back to Files API + importFile method for {filename}")

            try:
                # Upload to Files API with explicit MIME type
                uploaded_api_file = client.files.upload(
                    file=filepath,
                    config={
                        'mime_type': mime_type,  # Files API DOES accept mime_type
                        'display_name': filename
                    }
                )
                logger.info(f"File uploaded to Files API: {uploaded_api_file.name}")

                # Import to File Search Store (importFile has different config schema)
                import_config_fallback = {}

                # importFile accepts custom_metadata and chunking_config, but NOT display_name
                if upload_config.get('custom_metadata'):
                    import_config_fallback['custom_metadata'] = upload_config['custom_metadata']
                if upload_config.get('chunking_config'):
                    import_config_fallback['chunking_config'] = upload_config['chunking_config']

                operation = client.file_search_stores.import_file(
                    file_search_store_name=target_store_name,
                    file_name=uploaded_api_file.name,
                    config=import_config_fallback if import_config_fallback else None
                )
                logger.info(f"Importing {filename} into File Search Store via importFile")

            except Exception as fallback_error:
                logger.error(f"Fallback method also failed: {fallback_error}")
                raise fallback_error

        # Wait for operation to complete with INCREASED TIMEOUT - 120 segundos
        logger.info("Waiting for file import to complete")
        max_wait = 120  # 2 minutes timeout - Timeout de 120 segundos
        wait_time = 0
        while not operation.done and wait_time < max_wait:
            time.sleep(3)
            operation = client.operations.get(operation)
            wait_time += 3
            if wait_time % 15 == 0:  # Log progress every 15 seconds
                logger.info(f"Still waiting... ({wait_time}s elapsed)")

        if not operation.done:
            logger.error(f"File processing timeout after {max_wait}s")
            return jsonify({'error': f'File processing timeout after {max_wait} seconds. The file may still be processing in the background.'}), 500

        # Extract document ID from operation response
        document_id = None
        if hasattr(operation, 'response') and operation.response:
            document_id = getattr(operation.response, 'name', None)

        # Check for operation errors
        if hasattr(operation, 'error') and operation.error:
            logger.error(f"Upload operation failed: {operation.error.message}")
            return jsonify({'error': f'Upload failed: {operation.error.message}'}), 500

        # Track uploaded file
        file_info = {
            'filename': filename,
            'size': file_size,
            'mime_type': mime_type,  # Store MIME type for debugging
            'uploaded_at': time.strftime('%Y-%m-%d %H:%M:%S'),
            'custom_metadata': custom_metadata,
            'chunking_config': chunking_config if chunking_config.get('enabled') else None,
            'document_id': document_id
        }
        uploaded_files.append(file_info)

        # IMPORTANT: Save state to persistence - Guardar estado
        save_state()

        # Clean up local file
        os.remove(filepath)
        logger.info(f"File {filename} successfully uploaded and imported")

        return jsonify({
            'success': True,
            'message': f'File "{filename}" uploaded and processed successfully',
            'filename': filename,
            'file_size': file_size,
            'mime_type': mime_type,
            'store_name': target_store_name,
            'document_id': document_id,
            'uploaded_files': uploaded_files
        })

    except Exception as e:
        logger.error(f"Error uploading file: {str(e)}")
        # Clean up file if it exists
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
        # Clean up API file if uploaded in fallback
        if uploaded_api_file:
            try:
                client.files.delete(uploaded_api_file.name)
                logger.info(f"Cleaned up temporary file: {uploaded_api_file.name}")
            except:
                pass
        return jsonify({'error': f'Error uploading file: {str(e)}'}), 500

# ============================================
# IMPORT FROM URL - Importar desde URL
# ============================================

@app.route('/import-url', methods=['POST'])
def import_from_url():
    """Download a file from URL and import into File Search Store"""
    global file_search_store, uploaded_files

    data = request.json
    url = data.get('url', '').strip()
    store_name = data.get('store_name', '')
    custom_metadata = data.get('metadata', {})

    if not url:
        return jsonify({'error': 'URL is required'}), 400

    if not url.startswith(('http://', 'https://')):
        return jsonify({'error': 'URL must start with http:// or https://'}), 400

    # SSRF protection: block private/internal IPs
    from urllib.parse import urlparse
    import socket
    import ipaddress
    try:
        parsed = urlparse(url)
        hostname = parsed.hostname
        if not hostname:
            return jsonify({'error': 'Invalid URL'}), 400
        ip = socket.gethostbyname(hostname)
        addr = ipaddress.ip_address(ip)
        if addr.is_private or addr.is_loopback or addr.is_link_local:
            return jsonify({'error': 'URLs pointing to private/internal networks are not allowed'}), 400
    except (socket.gaierror, ValueError) as e:
        return jsonify({'error': f'Cannot resolve hostname: {str(e)}'}), 400

    filepath = None
    max_size = 100 * 1024 * 1024  # 100 MB
    try:
        # Download file from URL
        logger.info(f"Downloading file from URL: {url}")
        resp = http_requests.get(url, timeout=120, stream=True)
        resp.raise_for_status()

        # Determine filename
        filename = url.split('/')[-1].split('?')[0] or 'downloaded_file'
        filename = secure_filename(filename)
        if not filename or filename == '':
            filename = 'downloaded_file'

        # Save to temp with real-time size check
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        downloaded_size = 0
        with open(filepath, 'wb') as f:
            for chunk in resp.iter_content(chunk_size=8192):
                downloaded_size += len(chunk)
                if downloaded_size > max_size:
                    f.close()
                    os.remove(filepath)
                    return jsonify({'error': 'File exceeds 100 MB limit'}), 400
                f.write(chunk)

        file_size = os.path.getsize(filepath)
        mime_type = get_mime_type(filename)
        logger.info(f"Downloaded {filename} ({file_size} bytes, {mime_type})")

        # Determine target store
        target_store = file_search_store
        target_store_name = file_search_store.name if file_search_store else None

        if store_name:
            try:
                target_store = client.file_search_stores.get(name=store_name)
                target_store_name = store_name
            except Exception:
                return jsonify({'error': f'Store not found: {store_name}'}), 404

        if not target_store:
            return jsonify({'error': 'No active store. Create one first.'}), 400

        # Upload to File Search Store
        operation = client.file_search_stores.upload_to_file_search_store(
            file=filepath,
            file_search_store_name=target_store_name,
            config={
                'display_name': filename,
            }
        )

        # Wait for operation
        max_wait = 120
        wait_time = 0
        while not operation.done and wait_time < max_wait:
            time.sleep(5)
            wait_time += 5
            operation = client.operations.get(operation)
            if wait_time % 15 == 0:
                logger.info(f"Import from URL: still waiting... ({wait_time}s elapsed)")

        if not operation.done:
            return jsonify({'error': 'Upload timed out. The file may still be processing.'}), 408

        # Check for operation error
        if hasattr(operation, 'error') and operation.error:
            error_msg = getattr(operation.error, 'message', str(operation.error))
            logger.error(f"Import URL operation failed: {error_msg}")
            os.remove(filepath)
            return jsonify({'error': f'Import failed: {error_msg}'}), 500

        # Get document ID
        document_id = None
        if hasattr(operation, 'response') and operation.response:
            document_id = getattr(operation.response, 'name', None)

        # Build metadata dict
        metadata_dict = {}
        if isinstance(custom_metadata, list):
            for item in custom_metadata:
                metadata_dict[item.get('key', '')] = item.get('value', '')
        elif isinstance(custom_metadata, dict):
            metadata_dict = custom_metadata
        metadata_dict['source_url'] = url

        # Track file
        uploaded_files.append({
            'filename': filename,
            'size': file_size,
            'uploaded_at': time.strftime('%Y-%m-%d %H:%M:%S'),
            'custom_metadata': metadata_dict,
            'document_id': document_id,
            'source': 'url'
        })
        save_state()

        # Clean up
        os.remove(filepath)

        return jsonify({
            'success': True,
            'message': f'File "{filename}" imported from URL successfully',
            'filename': filename,
            'file_size': file_size,
            'mime_type': mime_type,
            'source_url': url,
            'document_id': document_id,
            'store_name': target_store_name
        })

    except http_requests.exceptions.RequestException as e:
        logger.error(f"Error downloading from URL: {str(e)}")
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'error': f'Error downloading file: {str(e)}'}), 400
    except Exception as e:
        logger.error(f"Error importing from URL: {str(e)}")
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'error': f'Error importing file: {str(e)}'}), 500

# ============================================
# CHAT WITH RAG & CITATIONS - Chat con RAG y citaciones
# Incluye MAX_HISTORY = 7 para límite de conversación
# ============================================

@app.route('/chat', methods=['POST'])
def chat():
    global conversation_history

    data = request.json
    user_message = data.get('message', '')
    metadata_filters = data.get('metadata_filters', [])  # Array de filtros del frontend
    system_prompt = data.get('system_prompt', '')
    structured_output = data.get('structured_output', False)
    response_schema = data.get('response_schema', None)
    top_k = data.get('top_k', None)  # Number of chunks File Search returns
    thinking_level = data.get('thinking_level', None)  # "high" or "low"
    media_resolution = data.get('media_resolution', None)  # "low", "medium", "high"

    # Model selection with whitelist
    ALLOWED_CHAT_MODELS = [
        'gemini-3.1-pro-preview',
        'gemini-3-flash-preview',
        'gemini-3.1-flash-lite-preview',
        'gemini-2.5-pro',
        'gemini-2.5-flash',
        'gemini-2.5-flash-lite',
    ]
    model = data.get('model', 'gemini-3-flash-preview')
    if model not in ALLOWED_CHAT_MODELS:
        model = 'gemini-3-flash-preview'

    if not user_message:
        return jsonify({'error': 'No message provided'}), 400

    if file_search_store is None:
        return jsonify({'error': 'Please upload a file first'}), 400

    try:
        # Build conversation context (last MAX_HISTORY messages) - Contexto de conversación
        # Note: user message is added to history AFTER successful API call
        context_messages = conversation_history[-MAX_HISTORY:]

        # Create prompt with conversation history
        prompt_parts = []

        # Add system prompt at the beginning if provided
        if system_prompt:
            prompt_parts.append(f"System Instructions: {system_prompt}\n")

        for msg in context_messages[:-1]:  # All except current message
            if msg['role'] == 'user':
                prompt_parts.append(f"User: {msg['content']}")
            else:
                prompt_parts.append(f"Assistant: {msg['content']}")

        # Add current question
        prompt_parts.append(f"User: {user_message}")
        prompt_parts.append("Assistant:")

        full_prompt = "\n\n".join(prompt_parts)

        logger.info(f"Querying with message: {user_message}")
        if metadata_filters:
            logger.info(f"Using metadata filters: {metadata_filters}")

        # Build file search config
        file_search_kwargs = {'file_search_store_names': [file_search_store.name]}
        if top_k is not None:
            try:
                file_search_kwargs['top_k'] = int(top_k)
                logger.info(f"Using top_k={top_k} for File Search")
            except (ValueError, TypeError):
                logger.warning(f"Invalid top_k value: {top_k}, ignoring")
        file_search_config = types.FileSearch(**file_search_kwargs)

        # Add metadata filters if provided - AIP-160 string format
        if metadata_filters and len(metadata_filters) > 0:
            # Build AIP-160 filter string: key=value pairs joined with AND
            # Docs: https://google.aip.dev/160
            filter_parts = []

            for filter_item in metadata_filters:
                key = filter_item.get('key', '')
                value = filter_item.get('value', '')

                if not key or not value:
                    continue

                # Determine if value is numeric or string
                try:
                    numeric_value = float(value)
                    # Numeric: no quotes
                    filter_parts.append(f'{key}={value}')
                except (ValueError, TypeError):
                    # String: wrap in quotes
                    escaped_value = str(value).replace('"', '\\"')
                    filter_parts.append(f'{key}="{escaped_value}"')

            if filter_parts:
                metadata_filter_string = " AND ".join(filter_parts)
                file_search_config.metadata_filter = metadata_filter_string
                logger.info(f"Applied metadata filter: {metadata_filter_string}")

        # Query with File Search
        logger.info(f"Chat using model: {model}")
        # Build generation config
        gen_config = {
            'tools': [types.Tool(file_search=file_search_config)]
        }

        # Add structured output if enabled (Gemini 3+ only)
        if structured_output and response_schema:
            gen_config['response_mime_type'] = 'application/json'
            gen_config['response_schema'] = response_schema
            logger.info("Structured output enabled with schema")

        # Add thinking_level if provided (Gemini 3+)
        if thinking_level in ('low', 'high'):
            gen_config['thinking_config'] = types.ThinkingConfig(thinking_budget=-1 if thinking_level == 'high' else 0)
            logger.info(f"Thinking level set to: {thinking_level}")

        # Add media_resolution if provided
        if media_resolution in ('low', 'medium', 'high'):
            gen_config['media_resolution'] = media_resolution
            logger.info(f"Media resolution set to: {media_resolution}")

        response = client.models.generate_content(
            model=model,
            contents=full_prompt,
            config=types.GenerateContentConfig(**gen_config)
        )

        assistant_message = response.text

        # Add both messages to history only after successful API call
        conversation_history.append({
            'role': 'user',
            'content': user_message
        })
        conversation_history.append({
            'role': 'assistant',
            'content': assistant_message
        })

        # Keep only last MAX_HISTORY messages
        if len(conversation_history) > MAX_HISTORY:
            conversation_history = conversation_history[-MAX_HISTORY:]

        # Extract grounding metadata (citations) - Extracción de citaciones
        metadata = None
        if response.candidates and len(response.candidates) > 0:
            candidate = response.candidates[0]
            if hasattr(candidate, 'grounding_metadata') and candidate.grounding_metadata:
                grounding = candidate.grounding_metadata

                # Extract citation information
                citations = []
                if hasattr(grounding, 'grounding_chunks') and grounding.grounding_chunks:
                    for chunk in grounding.grounding_chunks:
                        if hasattr(chunk, 'retrieved_context'):
                            ctx = chunk.retrieved_context
                            citation = {}
                            if hasattr(ctx, 'title'):
                                citation['title'] = ctx.title
                            if hasattr(ctx, 'uri'):
                                citation['uri'] = ctx.uri
                            if hasattr(ctx, 'text'):
                                citation['text'] = ctx.text
                            citations.append(citation)

                metadata = {
                    'citations': citations,
                    'citation_count': len(citations)
                }

        logger.info(f"Response generated successfully with {len(metadata['citations']) if metadata else 0} citations")

        return jsonify({
            'success': True,
            'response': assistant_message,
            'is_structured': bool(structured_output and response_schema),
            'metadata': metadata,
            'conversation_length': len(conversation_history),
            'model_used': model,
            'metadata_filters_applied': metadata_filters if metadata_filters else None,
            'filters_count': len(metadata_filters) if metadata_filters else 0
        })

    except Exception as e:
        logger.error(f"Error in chat: {str(e)}")
        return jsonify({'error': f'Error processing message: {str(e)}'}), 500

# ============================================
# FILE MANAGEMENT ENDPOINTS - Endpoints de gestión de archivos
# ============================================

@app.route('/delete-file/<int:file_index>', methods=['DELETE'])
def delete_file(file_index):
    global uploaded_files

    try:
        if file_index < 0 or file_index >= len(uploaded_files):
            return jsonify({'error': 'Invalid file index'}), 400

        file_info = uploaded_files[file_index]

        # Delete from Files API
        if file_info.get('file_api_name'):
            try:
                client.files.delete(file_info['file_api_name'])
                logger.info(f"Deleted file from Files API: {file_info['file_api_name']}")
            except Exception as e:
                logger.warning(f"Could not delete from Files API: {str(e)}")

        # Remove from tracking
        deleted_file = uploaded_files.pop(file_index)
        logger.info(f"Removed file from tracking: {deleted_file['filename']}")

        # Save state to persistence
        save_state()

        return jsonify({
            'success': True,
            'message': f"File '{deleted_file['filename']}' deleted successfully",
            'uploaded_files': uploaded_files
        })

    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        return jsonify({'error': f'Error deleting file: {str(e)}'}), 500

@app.route('/store-info', methods=['GET'])
def get_store_info():
    try:
        if file_search_store is None:
            return jsonify({
                'success': True,
                'store_exists': False,
                'message': 'No file search store created yet'
            })

        # Get store details
        store_details = client.file_search_stores.get(name=file_search_store.name)

        store_info = {
            'success': True,
            'store_exists': True,
            'name': store_details.name,
            'display_name': getattr(store_details, 'display_name', 'N/A'),
            'create_time': getattr(store_details, 'create_time', 'N/A'),
            'update_time': getattr(store_details, 'update_time', 'N/A'),
            'document_count': len(uploaded_files)
        }

        return jsonify(store_info)

    except Exception as e:
        logger.error(f"Error getting store info: {str(e)}")
        return jsonify({'error': f'Error getting store info: {str(e)}'}), 500

@app.route('/stores', methods=['GET'])
def list_stores():
    """List all File Search stores with their documents - Listar todos los stores con sus documentos"""
    try:
        stores = []
        for store in client.file_search_stores.list():
            # Get documents for this store
            documents = []
            try:
                for doc in client.file_search_stores.documents.list(parent=store.name):
                    custom_metadata = extract_document_metadata(doc, uploaded_files)

                    documents.append({
                        'name': doc.name,
                        'display_name': getattr(doc, 'display_name', 'N/A'),
                        'displayName': getattr(doc, 'display_name', 'N/A'),
                        'state': getattr(doc, 'state', 'UNKNOWN'),
                        'size_bytes': getattr(doc, 'size_bytes', 0),
                        'sizeBytes': getattr(doc, 'size_bytes', 0),
                        'mime_type': getattr(doc, 'mime_type', 'N/A'),
                        'mimeType': getattr(doc, 'mime_type', 'N/A'),
                        'create_time': str(getattr(doc, 'create_time', 'N/A')),
                        'createTime': str(getattr(doc, 'create_time', 'N/A')),
                        'custom_metadata': custom_metadata,
                        'customMetadata': custom_metadata
                    })
            except Exception as doc_error:
                logger.warning(f"Error listing documents for store {store.name}: {str(doc_error)}")

            store_size = int(getattr(store, 'size_bytes', 0) or 0)
            stores.append({
                'name': store.name,
                'display_name': getattr(store, 'display_name', 'N/A'),
                'displayName': getattr(store, 'display_name', 'N/A'),
                'create_time': str(getattr(store, 'create_time', 'N/A')),
                'createTime': str(getattr(store, 'create_time', 'N/A')),
                'size_bytes': store_size,
                'sizeBytes': store_size,
                'active_documents_count': getattr(store, 'active_documents_count', 0),
                'activeDocumentsCount': getattr(store, 'active_documents_count', 0),
                'pending_documents_count': getattr(store, 'pending_documents_count', 0),
                'pendingDocumentsCount': getattr(store, 'pending_documents_count', 0),
                'failed_documents_count': getattr(store, 'failed_documents_count', 0),
                'failedDocumentsCount': getattr(store, 'failed_documents_count', 0),
                'documents': documents
            })

        # Get current store name
        current_store_name = file_search_store.name if file_search_store else None

        return jsonify({
            'success': True,
            'stores': stores,
            'count': len(stores),
            'current_store': current_store_name
        })

    except Exception as e:
        logger.error(f"Error listing stores: {str(e)}")
        return jsonify({'error': f'Error listing stores: {str(e)}'}), 500

@app.route('/storage-usage', methods=['GET'])
def storage_usage():
    """Get aggregated storage usage across all stores - Uso de almacenamiento total"""
    try:
        total_size = 0
        store_count = 0
        for store in client.file_search_stores.list():
            total_size += int(getattr(store, 'size_bytes', 0) or 0)
            store_count += 1

        tier_limits = {
            'free': 1 * 1024**3,
            'tier1': 10 * 1024**3,
            'tier2': 100 * 1024**3,
            'tier3': 1 * 1024**4
        }

        current_tier = load_state_value('current_tier', 'free')
        tier_limit = tier_limits.get(current_tier, tier_limits['free'])
        usage_pct = (total_size / tier_limit * 100) if tier_limit > 0 else 0

        return jsonify({
            'success': True,
            'total_size_bytes': total_size,
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'store_count': store_count,
            'current_tier': current_tier,
            'tier_limit_bytes': tier_limit,
            'tier_limit_gb': round(tier_limit / (1024**3), 1),
            'usage_percentage': round(usage_pct, 2),
            'tier_limits': {k: round(v / (1024**3), 1) for k, v in tier_limits.items()}
        })
    except Exception as e:
        logger.error(f"Error getting storage usage: {str(e)}")
        return jsonify({'error': f'Error getting storage usage: {str(e)}'}), 500

@app.route('/update-tier', methods=['POST'])
def update_tier():
    """Update current tier setting - Actualizar tier actual"""
    data = request.json
    tier = data.get('tier', 'free')
    if tier not in ['free', 'tier1', 'tier2', 'tier3']:
        return jsonify({'error': 'Invalid tier'}), 400
    save_state_value('current_tier', tier)
    return jsonify({'success': True, 'tier': tier})

@app.route('/delete-store', methods=['DELETE'])
def delete_store():
    """Delete a File Search store (with all documents) - Eliminar store con todos los documentos"""
    global file_search_store, uploaded_files

    try:
        # Get store name from request body or use current store
        data = request.get_json() or {}
        store_name = data.get('store_name')

        # If no store name provided, use current store
        if not store_name:
            if file_search_store is None:
                return jsonify({'error': 'No store specified and no current store'}), 400
            store_name = file_search_store.name

        # Delete the store with force=True (deletes all documents automatically)
        client.file_search_stores.delete(name=store_name, config={'force': True})
        logger.info(f"Deleted file search store: {store_name}")

        # Reset state if we deleted the current store
        if file_search_store and file_search_store.name == store_name:
            file_search_store = None
            uploaded_files = []
            save_state()

        return jsonify({
            'success': True,
            'message': f'File search store deleted successfully: {store_name}'
        })

    except Exception as e:
        logger.error(f"Error deleting store: {str(e)}")
        return jsonify({'error': f'Error deleting store: {str(e)}'}), 500

# ============================================
# API DOCUMENTATION & INFO ENDPOINTS - Endpoints de documentación API
# ============================================

@app.route('/api-info', methods=['GET'])
def get_api_info():
    """Return API information for documentation tab - Devolver información API para pestaña de documentación"""
    try:
        api_info = {
            'success': True,
            'api_key': f"{api_key[:8]}...{api_key[-4:]}" if api_key and len(api_key) > 12 else '***',
            'store_exists': file_search_store is not None,
            'store_name': file_search_store.name if file_search_store else None,
            'store_display_name': getattr(file_search_store, 'display_name', 'RAG-App-Store') if file_search_store else 'RAG-App-Store',
            'file_count': len(uploaded_files),
            'files': uploaded_files,
            'model': 'gemini-3-flash-preview',
            'example_metadata_filters': []
        }

        # Collect example metadata keys from uploaded files
        metadata_keys = set()
        for file_info in uploaded_files:
            if file_info.get('custom_metadata'):
                for key in file_info['custom_metadata'].keys():
                    metadata_keys.add(key)

        api_info['metadata_keys'] = list(metadata_keys)

        return jsonify(api_info)

    except Exception as e:
        logger.error(f"Error getting API info: {str(e)}")
        return jsonify({'error': f'Error getting API info: {str(e)}'}), 500

@app.route('/update-api-key', methods=['POST'])
def update_api_key():
    """Update API key in .env file and reinitialize client - Actualizar API key"""
    global api_key, client
    try:
        data = request.json
        new_api_key = data.get('api_key', '').strip()

        if not new_api_key:
            return jsonify({'error': 'API key cannot be empty'}), 400

        # Update .env file
        env_path = '.env'
        if os.path.exists(env_path):
            with open(env_path, 'r') as f:
                lines = f.readlines()

            with open(env_path, 'w') as f:
                found = False
                for line in lines:
                    if line.startswith('GEMINI_API_KEY='):
                        f.write(f'GEMINI_API_KEY={new_api_key}\n')
                        found = True
                    else:
                        f.write(line)

                if not found:
                    f.write(f'\nGEMINI_API_KEY={new_api_key}\n')
        else:
            with open(env_path, 'w') as f:
                f.write(f'GEMINI_API_KEY={new_api_key}\n')

        # Update runtime variable and reinitialize client
        api_key = new_api_key
        client = genai.Client(api_key=api_key)

        logger.info("API key updated successfully")
        return jsonify({'success': True, 'message': 'API key updated successfully. Please reload the page.'})

    except Exception as e:
        logger.error(f"Error updating API key: {str(e)}")
        return jsonify({'error': f'Error updating API key: {str(e)}'}), 500

@app.route('/create-store', methods=['POST'])
def create_store():
    """Create a new File Search Store - Crear un nuevo File Search Store"""
    global file_search_store, uploaded_files
    try:
        data = request.json
        display_name = data.get('display_name', '').strip()

        if not display_name:
            return jsonify({'error': 'Display name is required'}), 400

        logger.info(f"Creating new File Search Store: {display_name}")

        # Create new store
        new_store = client.file_search_stores.create(
            config={'display_name': display_name}
        )

        logger.info(f"Created File Search Store: {new_store.name}")

        # Switch to the new store
        file_search_store = new_store
        uploaded_files = []

        # Save state
        save_state()

        return jsonify({
            'success': True,
            'store_name': new_store.name,
            'display_name': new_store.display_name,
            'message': f'Store "{display_name}" created successfully'
        })

    except Exception as e:
        logger.error(f"Error creating store: {str(e)}")
        return jsonify({'error': f'Error creating store: {str(e)}'}), 500

@app.route('/switch-store', methods=['POST'])
def switch_store():
    """Switch to a different File Search Store - Cambiar a otro store"""
    global file_search_store, uploaded_files
    try:
        data = request.json
        store_name = data.get('store_name', '').strip()

        if not store_name:
            return jsonify({'error': 'Store name is required'}), 400

        # Verify the store exists
        try:
            new_store = client.file_search_stores.get(name=store_name)
            file_search_store = new_store

            # Reset uploaded files list (will be loaded from store if needed)
            uploaded_files = []

            # Save new state
            save_state()

            logger.info(f"Switched to store: {store_name}")
            return jsonify({
                'success': True,
                'message': f'Switched to store: {store_name}',
                'store': {
                    'name': new_store.name,
                    'display_name': new_store.display_name or new_store.name.split('/')[-1],
                    'active_documents_count': new_store.active_documents_count or 0
                }
            })
        except Exception as e:
            return jsonify({'error': f'Store not found: {str(e)}'}), 404

    except Exception as e:
        logger.error(f"Error switching store: {str(e)}")
        return jsonify({'error': f'Error switching store: {str(e)}'}), 500

@app.route('/delete-document', methods=['DELETE'])
def delete_document():
    """Delete a specific document from a store - Eliminar un documento específico de un store"""
    global uploaded_files
    try:
        data = request.json
        document_name = data.get('document_name', '').strip()

        if not document_name:
            return jsonify({'error': 'Document name is required'}), 400

        # Delete the document with force=True
        try:
            client.file_search_stores.documents.delete(name=document_name, config={'force': True})
            logger.info(f"Deleted document: {document_name}")

            # If this document belongs to the current store, update uploaded_files
            if file_search_store and document_name.startswith(file_search_store.name):
                # Remove from uploaded_files if exists
                uploaded_files = [f for f in uploaded_files if f.get('document_id') != document_name]
                save_state()

            return jsonify({
                'success': True,
                'message': f'Document deleted successfully',
                'document_name': document_name
            })
        except Exception as e:
            return jsonify({'error': f'Document not found or could not be deleted: {str(e)}'}), 404

    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        return jsonify({'error': f'Error deleting document: {str(e)}'}), 500

@app.route('/update-document-metadata', methods=['POST'])
def update_document_metadata():
    """Update metadata for a document - Actualizar metadatos de un documento"""
    global uploaded_files
    try:
        data = request.json
        document_name = data.get('document_name', '').strip()
        new_metadata = data.get('metadata', {})

        if not document_name:
            return jsonify({'error': 'Document name is required'}), 400

        # Update metadata in local storage
        updated = False
        for file_info in uploaded_files:
            if file_info.get('document_id') == document_name:
                file_info['custom_metadata'] = new_metadata
                file_info['metadata_updated_at'] = time.strftime('%Y-%m-%d %H:%M:%S')
                updated = True
                break

        # If not found in uploaded_files, add new entry
        if not updated:
            uploaded_files.append({
                'document_id': document_name,
                'custom_metadata': new_metadata,
                'metadata_updated_at': time.strftime('%Y-%m-%d %H:%M:%S')
            })

        save_state()
        logger.info(f"Updated metadata for document: {document_name}")

        return jsonify({
            'success': True,
            'message': 'Metadata updated successfully',
            'document_name': document_name,
            'metadata': new_metadata
        })

    except Exception as e:
        logger.error(f"Error updating document metadata: {str(e)}")
        return jsonify({'error': f'Error updating metadata: {str(e)}'}), 500

@app.route('/suggest-metadata', methods=['POST'])
def suggest_metadata():
    """Analyze document with Gemini and suggest metadata - Sugerir metadatos con IA"""
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Get model preference (default to gemini-3-flash-preview)
        model = request.form.get('model', 'gemini-3-flash-preview')

        # Get language preference (default to 'en')
        language = request.form.get('language', 'en')
        logger.info(f"Using model for metadata generation: {model}, language: {language}")

        # Read file content
        file_content = file.read()
        file.seek(0)  # Reset file pointer for potential reuse

        # Determine mime type using our smart detection function
        mime_type = get_mime_type(file.filename)

        logger.info(f"Analyzing document: {file.filename} ({mime_type})")

        # Save file temporarily
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
        file.save(temp_path)

        # Determine if we need to extract text or use Files API
        use_text_extraction = mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # DOCX
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # XLSX
            'application/vnd.ms-excel',  # XLS
            'application/msword'  # DOC (old format)
        ]

        if use_text_extraction:
            # Extract text for unsupported formats (DOCX, XLSX)
            logger.info(f"Extracting text from {mime_type} file")

            if 'wordprocessingml' in mime_type or mime_type == 'application/msword':
                # DOCX or DOC
                extracted_text = extract_text_from_docx(temp_path)
            elif 'spreadsheetml' in mime_type or mime_type == 'application/vnd.ms-excel':
                # XLSX or XLS
                extracted_text = extract_text_from_xlsx(temp_path)
            else:
                extracted_text = ""

            if not extracted_text:
                raise Exception("Could not extract text from document")

            logger.info(f"Extracted {len(extracted_text)} characters")
            uploaded_file = None  # No uploaded file for text extraction
        else:
            # Upload file to Gemini Files API for supported formats (PDF, images, etc.)
            logger.info(f"Uploading to Files API: {mime_type}")
            uploaded_file = client.files.upload(
                file=temp_path,
                config={
                    'mime_type': mime_type,
                    'display_name': file.filename
                }
            )

        # System prompt with best practices - BILINGUAL VERSION
        if language == 'es':
            system_instruction = """Eres un asistente experto en análisis de documentos y gestión de metadatos para sistemas RAG (Retrieval Augmented Generation).

Tu trabajo es analizar documentos y extraer metadatos útiles que permitan:
1. Búsquedas eficientes por tipo, categoría, fecha, autor, etc.
2. Filtrado preciso en consultas RAG
3. Organización inteligente de información

FILOSOFÍA: CALIDAD sobre CANTIDAD
- Sugiere entre 5-12 campos de metadatos (máximo 20 solo si es absolutamente necesario)
- Cada campo de metadatos debe tener un propósito claro para búsqueda o filtrado
- Evita redundancias y campos triviales
- Si un campo no añade valor real para encontrar el documento, NO lo incluyas

REGLAS DE EXTRACCIÓN:
- Nombres de campos: minúsculas, sin espacios, usa guiones bajos
- Valores: strings simples o números, no arrays ni objetos
- Solo incluye metadatos que aparezcan explícitamente en el documento
- Prioriza metadatos que realmente diferencien este documento de otros
- IMPORTANTE: Todos los valores de metadatos deben estar EN ESPAÑOL

TIPOS DE DOCUMENTOS COMUNES:
- Facturas: numero_factura, proveedor, cliente, importe, fecha, estado
- Contratos: partes, fecha_firma, vigencia, valor, area, estado
- Informes: autor, departamento, periodo, tipo, confidencialidad
- Manuales: producto, version, idioma, categoria
- Emails: remitente, destinatario, asunto, fecha, prioridad

METADATOS UNIVERSALES (si aplican):
- titulo: título descriptivo del documento (extraído del contenido, NO del nombre de archivo)
- tipo: categoría del documento
- fecha: fecha principal (formato YYYY-MM-DD)
- autor: creador o responsable
- departamento: área organizacional
- confidencialidad: publico/interno/confidencial/privado
- estado: activo/pendiente/archivado/etc
- etiquetas: palabras clave (máximo 5, separadas por comas)

IMPORTANTE SOBRE EL TÍTULO:
- Debe extraerse del CONTENIDO del documento (encabezado, primera página, asunto)
- NO uses el nombre de archivo como título
- Debe ser descriptivo y profesional
- Formato: "Tipo de Documento - Descripción Breve"
- Ejemplo: "Factura de Servicios - Microsoft Azure Noviembre 2024"

Responde SOLO con JSON válido, sin markdown ni explicaciones."""
        else:
            system_instruction = """You are an expert assistant in document analysis and metadata management for RAG (Retrieval Augmented Generation) systems.

Your job is to analyze documents and extract useful metadata that enables:
1. Efficient searches by type, category, date, author, etc.
2. Precise filtering in RAG queries
3. Intelligent information organization

PHILOSOPHY: QUALITY over QUANTITY
- Suggest between 5-12 metadata fields (maximum 20 only if absolutely necessary)
- Each metadata field must have a clear purpose for search or filtering
- Avoid redundancies and trivial fields
- If a field doesn't add real value for finding the document, DON'T include it

EXTRACTION RULES:
- Field names: lowercase, no spaces, use underscores
- Values: simple strings or numbers, no arrays or objects
- Only include metadata that explicitly appears in the document
- Prioritize metadata that truly differentiates this document from others

COMMON DOCUMENT TYPES:
- Invoices: number, supplier, customer, amount, date, status
- Contracts: parties, signature_date, validity, value, area, status
- Reports: author, department, period, type, confidentiality
- Manuals: product, version, language, category
- Emails: sender, recipient, subject, date, priority

UNIVERSAL METADATA (if applicable):
- title: descriptive document title (extracted from content, NOT from filename)
- type: document category
- date: main date (format YYYY-MM-DD)
- author: creator or responsible party
- department: organizational area
- confidentiality: public/internal/confidential/private
- status: active/pending/archived/etc
- tags: keywords (maximum 5, comma-separated)

IMPORTANT ABOUT TITLE:
- Must be extracted from document CONTENT (header, first page, subject)
- DO NOT use the filename as title
- Must be descriptive and professional
- Format: "Document Type - Brief Description"
- Example: "Service Invoice - Microsoft Azure November 2024"

Respond ONLY with valid JSON, without markdown or explanations."""

        # Create analysis prompt - BILINGUAL VERSION
        if language == 'es':
            analysis_prompt = f"""Analiza este documento y extrae metadatos relevantes.

Nombre del archivo: {file.filename}

Proporciona un JSON con los metadatos encontrados. Estructura de ejemplo:
{{
    "titulo": "Factura de Servicios Cloud - Empresa XYZ",
    "tipo": "factura",
    "numero_factura": "INV-2024-001",
    "proveedor": "Empresa XYZ",
    "cliente": "Mi Empresa S.L.",
    "importe": "1500.00",
    "fecha": "2024-11-15",
    "estado": "pendiente"
}}

Recuerda:
- SIEMPRE incluye "titulo" extraído del CONTENIDO del documento
- CALIDAD sobre CANTIDAD: entre 5-12 campos de metadatos útiles (máximo 20)
- Solo campos que realmente ayuden a encontrar o filtrar este documento
- Evita campos obvios o redundantes
- Nombres de campos en minúsculas con guiones bajos
- TODOS los valores de metadatos en ESPAÑOL
- Formato JSON puro sin markdown"""
        else:
            analysis_prompt = f"""Analyze this document and extract relevant metadata.

Filename: {file.filename}

Provide a JSON with the metadata found. Example structure:
{{
    "title": "Cloud Services Invoice - XYZ Company",
    "type": "invoice",
    "invoice_number": "INV-2024-001",
    "supplier": "XYZ Company",
    "customer": "My Company Ltd",
    "amount": "1500.00",
    "date": "2024-11-15",
    "status": "pending"
}}

Remember:
- ALWAYS include "title" extracted from document CONTENT
- QUALITY over QUANTITY: between 5-12 useful metadata fields (maximum 20)
- Only fields that really help find or filter this document
- Avoid obvious or redundant fields
- Field names in lowercase with underscores
- ALL metadata values in ENGLISH
- Pure JSON format without markdown"""

        # Generate content with system instruction using selected model
        if use_text_extraction:
            # Use extracted text for DOCX/XLSX
            logger.info("Generating metadata from extracted text")
            response = client.models.generate_content(
                model=model,
                contents=[
                    types.Content(
                        role='user',
                        parts=[
                            types.Part.from_text(text=f"{analysis_prompt}\n\nDocument content:\n{extracted_text[:10000]}")  # Limit to 10k chars
                        ]
                    )
                ],
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.1
                )
            )
        else:
            # Use Files API for PDF and other supported formats
            logger.info("Generating metadata from uploaded file")
            response = client.models.generate_content(
                model=model,
                contents=[
                    types.Content(
                        role='user',
                        parts=[
                            types.Part.from_uri(
                                file_uri=uploaded_file.uri,
                                mime_type=uploaded_file.mime_type
                            ),
                            types.Part.from_text(text=analysis_prompt)
                        ]
                    )
                ],
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.1
                )
            )

        # Extract JSON from response
        response_text = response.text.strip()

        # Clean response - remove markdown code blocks if present
        if response_text.startswith('```'):
            # Remove ```json and ``` markers
            response_text = response_text.replace('```json', '').replace('```', '').strip()

        # Parse JSON
        suggested_metadata = json.loads(response_text)

        logger.info(f"Suggested metadata: {suggested_metadata}")

        # Clean up temporary files
        try:
            # Only delete from Files API if we uploaded there
            if uploaded_file is not None:
                client.files.delete(name=uploaded_file.name)

            # Always delete local temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except Exception as cleanup_error:
            logger.warning(f"Could not delete temporary file: {cleanup_error}")

        return jsonify({
            'success': True,
            'metadata': suggested_metadata,
            'filename': file.filename
        })

    except json.JSONDecodeError as json_error:
        logger.error(f"Error parsing JSON from Gemini: {json_error}")
        logger.error(f"Response was: {response_text}")
        return jsonify({
            'error': 'Error parsing AI response',
            'details': f'Could not parse JSON: {str(json_error)}'
        }), 500
    except Exception as e:
        logger.error(f"Error suggesting metadata: {str(e)}")
        return jsonify({'error': f'Error analyzing document: {str(e)}'}), 500

# ============================================
# DOCUMENT QUERY - Búsqueda semántica en documento específico
# ============================================

@app.route('/document-query', methods=['POST'])
def document_query():
    """Perform semantic search on a specific document WITHOUT full generation.

    Uses the documents.query API to retrieve relevant chunks directly.
    Useful for retrieval-only use cases or inspecting what chunks match a query.

    Request body:
        document_name (str): Full document resource name
            e.g. "fileSearchStores/.../documents/..."
        query (str): The semantic search query
        results_count (int): Number of chunks to return, max 100 (default 10)
        metadata_filters (list): Optional AIP-160 metadata filters

    Returns:
        JSON with list of relevant chunks and their content
    """
    try:
        data = request.json
        document_name = data.get('document_name', '').strip()
        query = data.get('query', '').strip()
        results_count = data.get('results_count', 10)
        metadata_filters = data.get('metadata_filters', [])

        if not document_name:
            return jsonify({'error': 'document_name is required'}), 400
        if not query:
            return jsonify({'error': 'query is required'}), 400

        # Clamp results_count to valid range
        try:
            results_count = max(1, min(100, int(results_count)))
        except (ValueError, TypeError):
            results_count = 10

        logger.info(f"Document query on {document_name}: '{query}' (results_count={results_count})")

        # Build query config
        query_config = {'results_count': results_count}

        # Add metadata filters if provided (AIP-160 format)
        if metadata_filters and len(metadata_filters) > 0:
            filter_parts = []
            for filter_item in metadata_filters:
                key = filter_item.get('key', '')
                value = filter_item.get('value', '')
                if not key or not value:
                    continue
                try:
                    float(value)
                    filter_parts.append(f'{key}={value}')
                except (ValueError, TypeError):
                    escaped_value = str(value).replace('"', '\\"')
                    filter_parts.append(f'{key}="{escaped_value}"')
            if filter_parts:
                query_config['metadata_filter'] = ' AND '.join(filter_parts)
                logger.info(f"Document query filter: {query_config['metadata_filter']}")

        # Execute semantic search via documents.query
        query_response = client.file_search_stores.documents.query(
            name=document_name,
            query=query,
            config=query_config
        )

        # Extract chunks from response
        chunks = []
        if hasattr(query_response, 'relevant_chunks') and query_response.relevant_chunks:
            for chunk in query_response.relevant_chunks:
                chunk_data = {}
                if hasattr(chunk, 'chunk_relevance_score'):
                    chunk_data['relevance_score'] = chunk.chunk_relevance_score
                if hasattr(chunk, 'chunk') and chunk.chunk:
                    c = chunk.chunk
                    if hasattr(c, 'data') and c.data:
                        chunk_data['text'] = getattr(c.data, 'string_value', '')
                    if hasattr(c, 'custom_metadata') and c.custom_metadata:
                        chunk_meta = {}
                        for m in c.custom_metadata:
                            k = getattr(m, 'key', '')
                            if hasattr(m, 'string_value'):
                                chunk_meta[k] = getattr(m, 'string_value', '')
                            elif hasattr(m, 'numeric_value'):
                                chunk_meta[k] = getattr(m, 'numeric_value', 0)
                        chunk_data['metadata'] = chunk_meta
                chunks.append(chunk_data)

        logger.info(f"Document query returned {len(chunks)} chunks")

        return jsonify({
            'success': True,
            'document_name': document_name,
            'query': query,
            'results_count': results_count,
            'chunks': chunks,
            'chunks_returned': len(chunks)
        })

    except Exception as e:
        logger.error(f"Error in document query: {str(e)}")
        return jsonify({'error': f'Error performing document query: {str(e)}'}), 500


# ============================================
# AUTO-ENRICHMENT ENDPOINT - Enriquecimiento automatico con schema estructurado
# ============================================

# Default enrichment schema for FixMe Malaga use case (tech repair conversations)
DEFAULT_ENRICH_SCHEMA = {
    'type': 'object',
    'properties': {
        'categoria_principal': {
            'type': 'string',
            'description': 'Main category: reparacion_pantalla, bateria, placa_base, software, consulta_precio, garantia, envio, reacondicionado, otro'
        },
        'categorias': {
            'type': 'array',
            'items': {'type': 'string'},
            'description': 'All categories that apply to this document'
        },
        'dispositivo': {
            'type': 'string',
            'description': 'Device model mentioned, e.g. iPhone 14, Samsung S21, HP Omen'
        },
        'marca': {
            'type': 'string',
            'description': 'Brand: Apple, Samsung, Xiaomi, HP, Lenovo, otro'
        },
        'estado_final': {
            'type': 'string',
            'description': 'Final status: reparado, pendiente, presupuesto_rechazado, no_reparable, en_proceso'
        },
        'sentimiento_cliente': {
            'type': 'string',
            'description': 'Client sentiment: positivo, negativo, neutro, mixto'
        },
        'resumen': {
            'type': 'string',
            'description': 'Brief 1-2 sentence summary of the document content'
        },
        'precio_mencionado': {
            'type': 'number',
            'description': 'Price mentioned in euros, 0 if none'
        },
        'es_recurrente': {
            'type': 'boolean',
            'description': 'Whether the client appears to be a recurring customer'
        }
    },
    'required': ['categoria_principal', 'categorias', 'marca', 'estado_final', 'sentimiento_cliente', 'resumen']
}


def _analyze_file_with_schema(file_path, filename, mime_type, schema, model='gemini-3-flash-preview'):
    """
    Core analysis: upload file to Gemini and extract structured metadata using response_schema.
    Handles DOCX/XLSX text extraction. Cleans up Files API upload in all cases.
    """
    use_text_extraction = mime_type in [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel',
        'application/msword'
    ]

    uploaded_file = None

    if use_text_extraction:
        logger.info(f"Auto-enrich: extracting text from {mime_type}")
        if 'wordprocessingml' in mime_type or mime_type == 'application/msword':
            extracted_text = extract_text_from_docx(file_path)
        elif 'spreadsheetml' in mime_type or mime_type == 'application/vnd.ms-excel':
            extracted_text = extract_text_from_xlsx(file_path)
        else:
            extracted_text = ''
        if not extracted_text:
            raise ValueError("Could not extract text from document")
    else:
        logger.info(f"Auto-enrich: uploading to Files API ({mime_type})")
        uploaded_file = client.files.upload(
            file=file_path,
            config={'mime_type': mime_type, 'display_name': filename}
        )

    system_instruction = (
        "You are a document analysis expert. Extract structured metadata from the document "
        "following the provided JSON schema exactly. Return only valid JSON with the requested fields. "
        "Base all values strictly on what appears in the document."
    )

    prompt = f"Analyze this document and extract structured metadata. Filename: {filename}"

    try:
        if use_text_extraction:
            content_parts = [
                types.Part.from_text(text=f"{prompt}\n\nDocument content:\n{extracted_text[:12000]}")
            ]
        else:
            content_parts = [
                types.Part.from_uri(
                    file_uri=uploaded_file.uri,
                    mime_type=uploaded_file.mime_type
                ),
                types.Part.from_text(text=prompt)
            ]

        response = client.models.generate_content(
            model=model,
            contents=[types.Content(role='user', parts=content_parts)],
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.1,
                response_mime_type='application/json',
                response_schema=schema
            )
        )

        return json.loads(response.text)

    finally:
        if uploaded_file is not None:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception as cleanup_err:
                logger.warning(f"Could not delete temp file from Files API: {cleanup_err}")


@app.route('/auto-enrich', methods=['POST'])
def auto_enrich():
    """
    Analyze a single document with structured output (guaranteed JSON schema).
    Accepts multipart/form-data with 'file' field.
    Optional form fields: 'schema' (JSON string), 'model'.
    Returns enriched metadata matching the schema.
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        model = request.form.get('model', 'gemini-3-flash-preview')

        custom_schema_raw = request.form.get('schema')
        if custom_schema_raw:
            try:
                schema = json.loads(custom_schema_raw)
            except json.JSONDecodeError:
                return jsonify({'error': 'Invalid JSON in schema parameter'}), 400
        else:
            schema = DEFAULT_ENRICH_SCHEMA

        mime_type = get_mime_type(file.filename)
        logger.info(f"Auto-enrich: {file.filename} ({mime_type}), model: {model}")

        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
        file.save(temp_path)

        try:
            enriched = _analyze_file_with_schema(temp_path, file.filename, mime_type, schema, model)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        logger.info(f"Auto-enrich result for {file.filename}: {enriched}")

        return jsonify({
            'success': True,
            'metadata': enriched,
            'filename': file.filename
        })

    except Exception as e:
        logger.error(f"Error in auto-enrich: {str(e)}")
        return jsonify({'error': f'Error enriching document: {str(e)}'}), 500


# ============================================
# UTILITY ENDPOINTS - Endpoints de utilidad
# ============================================

@app.route('/clear', methods=['POST'])
def clear_conversation():
    global conversation_history
    conversation_history = []
    logger.info("Conversation history cleared")
    return jsonify({'success': True, 'message': 'Conversation cleared'})

@app.route('/files', methods=['GET'])
def get_files():
    return jsonify({
        'success': True,
        'files': uploaded_files,
        'store_name': file_search_store.name if file_search_store else None
    })

@app.route('/current-store-documents', methods=['GET'])
def get_current_store_documents():
    """Get all documents from the current store with metadata"""
    try:
        if not file_search_store:
            return jsonify({
                'success': True,
                'documents': [],
                'store_name': None,
                'message': 'No store selected'
            })

        documents = []
        try:
            for doc in client.file_search_stores.documents.list(parent=file_search_store.name):
                custom_metadata = extract_document_metadata(doc, uploaded_files)

                documents.append({
                    'name': doc.name,
                    'display_name': getattr(doc, 'display_name', 'N/A'),
                    'displayName': getattr(doc, 'display_name', 'N/A'),
                    'state': getattr(doc, 'state', 'UNKNOWN'),
                    'size_bytes': getattr(doc, 'size_bytes', 0),
                    'sizeBytes': getattr(doc, 'size_bytes', 0),
                    'mime_type': getattr(doc, 'mime_type', 'N/A'),
                    'mimeType': getattr(doc, 'mime_type', 'N/A'),
                    'create_time': str(getattr(doc, 'create_time', 'N/A')),
                    'createTime': str(getattr(doc, 'create_time', 'N/A')),
                    'custom_metadata': custom_metadata,
                    'customMetadata': custom_metadata
                })
        except Exception as doc_error:
            logger.warning(f"Error listing documents for current store: {str(doc_error)}")

        return jsonify({
            'success': True,
            'documents': documents,
            'store_name': file_search_store.name,
            'store_display_name': getattr(file_search_store, 'display_name', 'N/A')
        })

    except Exception as e:
        logger.error(f"Error getting current store documents: {str(e)}")
        return jsonify({'error': f'Error getting documents: {str(e)}'}), 500

@app.route('/status', methods=['GET'])
def status():
    return jsonify({
        'file_uploaded': file_search_store is not None,
        'conversation_length': len(conversation_history),
        'store_name': file_search_store.name if file_search_store else None,
        'uploaded_files': uploaded_files
    })

# ============================================
# INVESTIGATIONS / REPORTS - Investigaciones multi-pregunta con RAG
# ============================================

@app.route('/investigate', methods=['POST'])
def investigate():
    """Execute a multi-question investigation using RAG with File Search.

    Each question is answered independently using semantic search on the store,
    then a final executive summary is generated from all answers.

    Request body:
        title (str): Title of the investigation
        questions (list[str]): List of questions to investigate
        store_name (str, optional): Store to use; defaults to current active store
        model (str, optional): Model to use; defaults to gemini-3.1-pro-preview

    Returns:
        JSON with the full investigation object including sections and summary
    """
    import uuid
    import datetime

    try:
        data = request.json
        title = data.get('title', '').strip()
        questions = [q.strip() for q in data.get('questions', []) if q.strip()]
        store_name = data.get('store_name', '')
        ALLOWED_MODELS = [
            'gemini-3.1-pro-preview', 'gemini-3-flash-preview', 'gemini-3.1-flash-lite-preview',
            'gemini-2.5-pro', 'gemini-2.5-flash', 'gemini-2.5-flash-lite'
        ]
        model = data.get('model', 'gemini-3.1-pro-preview')
        if model not in ALLOWED_MODELS:
            model = 'gemini-3.1-pro-preview'

        if not title:
            return jsonify({'error': 'Title is required'}), 400
        if not questions or len(questions) == 0:
            return jsonify({'error': 'At least one question is required'}), 400

        # Determine target store
        target_store = store_name or (file_search_store.name if file_search_store else None)
        if not target_store:
            return jsonify({'error': 'No active store. Create one and upload documents first.'}), 400

        logger.info(f"Starting investigation '{title}' with {len(questions)} questions on store {target_store}")

        sections = []
        for i, question in enumerate(questions):
            logger.info(f"Processing question {i+1}/{len(questions)}: {question[:80]}...")
            try:
                response = client.models.generate_content(
                    model=model,
                    contents=question,
                    config=types.GenerateContentConfig(
                        tools=[types.Tool(
                            file_search=types.FileSearch(
                                file_search_store_names=[target_store]
                            )
                        )]
                    )
                )

                # Extract citations from grounding metadata
                citations = []
                if response.candidates and response.candidates[0].grounding_metadata:
                    grounding = response.candidates[0].grounding_metadata
                    if hasattr(grounding, 'grounding_chunks') and grounding.grounding_chunks:
                        for chunk in grounding.grounding_chunks:
                            if hasattr(chunk, 'retrieved_context'):
                                ctx = chunk.retrieved_context
                                citations.append({
                                    'title': getattr(ctx, 'title', ''),
                                    'text': getattr(ctx, 'text', '')[:200]
                                })

                sections.append({
                    'question': question,
                    'answer': response.text,
                    'citations': citations
                })

            except Exception as q_error:
                logger.error(f"Error processing question {i+1}: {str(q_error)}")
                sections.append({
                    'question': question,
                    'answer': f'Error al procesar esta pregunta: {str(q_error)}',
                    'citations': []
                })

        # Generate executive summary from all answers
        logger.info("Generating executive summary...")
        all_answers = "\n\n".join([
            f"Pregunta: {s['question']}\nRespuesta: {s['answer']}"
            for s in sections
        ])
        try:
            summary_response = client.models.generate_content(
                model=model,
                contents=(
                    f"Eres un analista experto. Genera un resumen ejecutivo conciso de 3-5 lineas "
                    f"de esta investigacion titulada '{title}':\n\n{all_answers}"
                )
            )
            summary = summary_response.text
        except Exception as summary_error:
            logger.error(f"Error generating summary: {str(summary_error)}")
            summary = "No se pudo generar el resumen ejecutivo."

        investigation = {
            'id': str(uuid.uuid4()),
            'title': title,
            'store_name': target_store,
            'sections': sections,
            'summary': summary,
            'created_at': datetime.datetime.now().isoformat(),
            'metadata': {
                'total_questions': len(questions),
                'total_citations': sum(len(s['citations']) for s in sections),
                'model_used': model
            }
        }

        # Persist to state file
        investigations = load_state_value('investigations', [])
        investigations.append(investigation)
        save_state_value('investigations', investigations)

        logger.info(f"Investigation '{title}' completed. ID: {investigation['id']}")

        return jsonify({'success': True, 'investigation': investigation})

    except Exception as e:
        logger.error(f"Error in investigate: {str(e)}")
        return jsonify({'error': f'Error running investigation: {str(e)}'}), 500


@app.route('/investigations', methods=['GET'])
def list_investigations():
    """List all saved investigations, sorted by creation date descending."""
    try:
        investigations = load_state_value('investigations', [])
        # Return newest first
        investigations_sorted = sorted(
            investigations,
            key=lambda x: x.get('created_at', ''),
            reverse=True
        )
        return jsonify({
            'success': True,
            'investigations': investigations_sorted,
            'count': len(investigations_sorted)
        })
    except Exception as e:
        logger.error(f"Error listing investigations: {str(e)}")
        return jsonify({'error': f'Error listing investigations: {str(e)}'}), 500


@app.route('/investigations/<investigation_id>', methods=['GET'])
def get_investigation(investigation_id):
    """Get a specific investigation by ID."""
    try:
        investigations = load_state_value('investigations', [])
        for inv in investigations:
            if inv.get('id') == investigation_id:
                return jsonify({'success': True, 'investigation': inv})
        return jsonify({'error': f'Investigation not found: {investigation_id}'}), 404
    except Exception as e:
        logger.error(f"Error getting investigation {investigation_id}: {str(e)}")
        return jsonify({'error': f'Error getting investigation: {str(e)}'}), 500


@app.route('/investigations/<investigation_id>', methods=['DELETE'])
def delete_investigation(investigation_id):
    """Delete a specific investigation by ID."""
    try:
        investigations = load_state_value('investigations', [])
        original_count = len(investigations)
        investigations = [inv for inv in investigations if inv.get('id') != investigation_id]

        if len(investigations) == original_count:
            return jsonify({'error': f'Investigation not found: {investigation_id}'}), 404

        save_state_value('investigations', investigations)
        logger.info(f"Deleted investigation: {investigation_id}")

        return jsonify({
            'success': True,
            'message': 'Investigation deleted successfully',
            'id': investigation_id
        })
    except Exception as e:
        logger.error(f"Error deleting investigation {investigation_id}: {str(e)}")
        return jsonify({'error': f'Error deleting investigation: {str(e)}'}), 500


@app.route('/investigations/<investigation_id>/export-pdf', methods=['POST'])
def export_investigation_pdf(investigation_id):
    """Export an investigation as a PDF report.

    Generates a formatted PDF with header/footer images, executive summary,
    numbered sections with questions/answers/citations.

    Returns:
        PDF file as attachment
    """
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, Image, Table, TableStyle)
    from reportlab.lib import colors
    from flask import send_file
    import re

    try:
        investigations = load_state_value('investigations', [])
        inv = next((i for i in investigations if i.get('id') == investigation_id), None)
        if not inv:
            return jsonify({'error': 'Investigation not found'}), 404

        # Corporate color palette
        AZUL = HexColor('#2B9FD1')
        AZUL_OSCURO = HexColor('#1A7BA8')
        GRIS_TEXTO = HexColor('#374151')
        GRIS_CLARO = HexColor('#6B7280')
        AZUL_FONDO = HexColor('#EFF6FF')
        AZUL_BORDE = HexColor('#BFDBFE')
        NEGRO = HexColor('#111827')

        # Header/footer image paths
        HEADER_IMG = os.path.expanduser('~/.claude/assets/webcomunica/cabecera.png')
        FOOTER_IMG = os.path.expanduser('~/.claude/assets/webcomunica/pie.png')

        PAGE_W, PAGE_H = A4
        MARGIN = 2 * cm

        buf = io.BytesIO()

        # Helper to strip markdown/HTML for plain text in PDF
        def strip_markup(text):
            if not text:
                return ''
            # Remove markdown headers
            text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
            # Remove bold/italic
            text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
            text = re.sub(r'_{1,3}(.*?)_{1,3}', r'\1', text)
            # Remove inline code
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # Remove links [text](url)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', text)
            # Escape reportlab special chars
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return text.strip()

        def make_doc():
            """Build the platypus story (list of flowables)."""
            styles = getSampleStyleSheet()

            style_title = ParagraphStyle(
                'InvTitle',
                fontName='Helvetica-Bold',
                fontSize=18,
                textColor=AZUL_OSCURO,
                spaceAfter=6,
                leading=22,
            )
            style_meta = ParagraphStyle(
                'InvMeta',
                fontName='Helvetica',
                fontSize=9,
                textColor=GRIS_CLARO,
                spaceAfter=4,
            )
            style_section_label = ParagraphStyle(
                'SectionLabel',
                fontName='Helvetica-Bold',
                fontSize=9,
                textColor=AZUL,
                spaceBefore=4,
                spaceAfter=2,
                textTransform='uppercase',
            )
            style_question = ParagraphStyle(
                'Question',
                fontName='Helvetica-Bold',
                fontSize=11,
                textColor=NEGRO,
                spaceBefore=4,
                spaceAfter=6,
            )
            style_answer = ParagraphStyle(
                'Answer',
                fontName='Helvetica',
                fontSize=10,
                textColor=GRIS_TEXTO,
                leading=15,
                spaceAfter=8,
                alignment=TA_JUSTIFY,
            )
            style_citation = ParagraphStyle(
                'Citation',
                fontName='Helvetica-Oblique',
                fontSize=8,
                textColor=AZUL_OSCURO,
                leftIndent=12,
                spaceAfter=2,
            )
            style_summary = ParagraphStyle(
                'Summary',
                fontName='Helvetica',
                fontSize=10,
                textColor=HexColor('#1E3A5F'),
                leading=15,
                spaceAfter=4,
                alignment=TA_JUSTIFY,
            )
            style_footer = ParagraphStyle(
                'Footer',
                fontName='Helvetica-Oblique',
                fontSize=8,
                textColor=GRIS_CLARO,
                alignment=TA_CENTER,
            )

            story = []

            # --- Header image ---
            if os.path.exists(HEADER_IMG):
                img = Image(HEADER_IMG, width=PAGE_W - 2 * MARGIN, height=2.2 * cm)
                story.append(img)
                story.append(Spacer(1, 0.4 * cm))

            # --- Title block ---
            title_text = strip_markup(inv.get('title', 'Investigacion'))
            story.append(Paragraph(title_text, style_title))

            # Metadata line
            created_at = inv.get('created_at', '')
            model_used = inv.get('metadata', {}).get('model_used', '')
            total_questions = inv.get('metadata', {}).get('total_questions', len(inv.get('sections', [])))
            total_citations = inv.get('metadata', {}).get('total_citations', 0)
            meta_parts = []
            if created_at:
                meta_parts.append(f"Fecha: {created_at[:19].replace('T', ' ')}")
            if model_used:
                meta_parts.append(f"Modelo: {model_used}")
            meta_parts.append(f"Preguntas: {total_questions}")
            meta_parts.append(f"Citas: {total_citations}")
            story.append(Paragraph(' &nbsp;|&nbsp; '.join(meta_parts), style_meta))
            story.append(HRFlowable(width='100%', thickness=1.5, color=AZUL, spaceAfter=10))

            # --- Executive summary box ---
            summary_text = strip_markup(inv.get('summary', ''))
            if summary_text:
                story.append(Paragraph('Resumen Ejecutivo', style_section_label))
                # Simulate the blue box with a Table of one cell
                summary_table = Table(
                    [[Paragraph(summary_text, style_summary)]],
                    colWidths=[PAGE_W - 2 * MARGIN],
                )
                summary_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), AZUL_FONDO),
                    ('BOX', (0, 0), (-1, -1), 1.5, AZUL_BORDE),
                    ('LEFTPADDING', (0, 0), (-1, -1), 12),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ('ROUNDEDCORNERS', [6, 6, 6, 6]),
                ]))
                story.append(summary_table)
                story.append(Spacer(1, 0.5 * cm))

            # --- Sections ---
            story.append(Paragraph('Preguntas y Respuestas', style_section_label))
            story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#E5E7EB'), spaceAfter=8))

            for idx, section in enumerate(inv.get('sections', []), start=1):
                question = strip_markup(section.get('question', ''))
                answer = strip_markup(section.get('answer', ''))
                citations = section.get('citations', [])

                story.append(Paragraph(f"{idx}. {question}", style_question))

                if answer:
                    # Split by newlines to preserve paragraph breaks
                    for para in answer.split('\n'):
                        para = para.strip()
                        if para:
                            story.append(Paragraph(para, style_answer))

                if citations:
                    story.append(Paragraph(
                        '<font color="#6B7280"><b>Fuentes:</b></font>', style_citation))
                    for c in citations:
                        cite_title = strip_markup(c.get('title', 'Documento'))
                        cite_text = strip_markup((c.get('text', '') or '')[:200])
                        if cite_text:
                            story.append(Paragraph(
                                f'&bull; <b>{cite_title}</b>: {cite_text}...', style_citation))
                        else:
                            story.append(Paragraph(f'&bull; <b>{cite_title}</b>', style_citation))

                story.append(HRFlowable(
                    width='100%', thickness=0.5, color=HexColor('#F3F4F6'), spaceAfter=6))

            # --- Footer image ---
            story.append(Spacer(1, 0.6 * cm))
            if os.path.exists(FOOTER_IMG):
                story.append(Image(FOOTER_IMG, width=PAGE_W - 2 * MARGIN, height=1.5 * cm))
            else:
                story.append(HRFlowable(width='100%', thickness=1, color=AZUL, spaceBefore=4))
                story.append(Spacer(1, 0.3 * cm))
                story.append(Paragraph(
                    'Generado con Gemini File Search Manager by Webcomunica',
                    style_footer))

            return story

        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=MARGIN,
            rightMargin=MARGIN,
            topMargin=MARGIN,
            bottomMargin=MARGIN,
        )
        doc.build(make_doc())
        buf.seek(0)

        short_id = investigation_id[:8]
        logger.info(f"PDF exported for investigation {investigation_id}")
        return send_file(
            buf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'investigacion_{short_id}.pdf',
        )

    except Exception as e:
        logger.error(f"Error exporting PDF for investigation {investigation_id}: {str(e)}")
        return jsonify({'error': f'Error generating PDF: {str(e)}'}), 500


# ============================================
# ============================================
# TTS - Text to Speech
# ============================================

@app.route('/tts', methods=['POST'])
def text_to_speech():
    """Convert text to speech using Gemini TTS"""
    import base64
    import struct

    data = request.json
    text = data.get('text', '')[:4000]  # Limit 4000 chars
    voice = data.get('voice', 'Aoede')  # Aoede, Charon, Fenrir, Kore, Puck, Leda

    if not text:
        return jsonify({'error': 'No text provided'}), 400

    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash-preview-tts',
            contents=text,
            config=types.GenerateContentConfig(
                response_modalities=['AUDIO'],
                speech_config=types.SpeechConfig(
                    voice_config=types.VoiceConfig(
                        prebuilt_voice_config=types.PrebuiltVoiceConfig(voice_name=voice)
                    )
                )
            )
        )

        # Extract PCM audio data (may be base64 string or bytes)
        raw_audio = response.candidates[0].content.parts[0].inline_data.data
        if isinstance(raw_audio, str):
            import base64 as _b64
            audio_data = _b64.b64decode(raw_audio)
        else:
            audio_data = raw_audio

        sample_rate = 24000  # Gemini TTS outputs 24kHz mono 16-bit

        # Convert PCM to WAV (mono, 16-bit)
        wav_header = struct.pack('<4sI4s4sIHHIIHH4sI',
            b'RIFF', 36 + len(audio_data), b'WAVE',
            b'fmt ', 16, 1, 1, sample_rate, sample_rate * 2, 2, 16,
            b'data', len(audio_data)
        )
        wav_bytes = wav_header + audio_data
        audio_base64 = base64.b64encode(wav_bytes).decode('utf-8')

        return jsonify({
            'success': True,
            'audio': audio_base64,
            'format': 'wav',
            'voice': voice
        })

    except Exception as e:
        logger.error(f"TTS error: {str(e)}")
        return jsonify({'error': f'TTS error: {str(e)}'}), 500

# ============================================
# APPLICATION ENTRY POINT - Punto de entrada de la aplicación
# Puerto configurado en 5001
# ============================================

if __name__ == '__main__':
    # Load persisted state on startup - Cargar estado persistido al iniciar
    load_state()
    debug_mode = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'
    app.run(debug=debug_mode, host='localhost', port=5001)
